from docx import Document
from datetime import datetime

def create_docx_report(title: str, exec_summary: str, eda_summary: dict,
                       model_summary: str, figs):
    """Create MS Word report."""
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(exec_summary)

    doc.add_heading("EDA Summary", level=1)
    for k, v in eda_summary.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Model Summary", level=1)
    doc.add_paragraph(model_summary)

    # Optionally embed figures
    for fig_path in figs:
        doc.add_picture(fig_path, width=None)

    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename
